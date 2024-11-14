import sys
from PyQt6.QtWidgets import QApplication, QMainWindow, QTextEdit, QVBoxLayout, QWidget, QFileDialog, QMessageBox, QColorDialog
from PyQt6.QtGui import QTextCharFormat, QFont, QColor
from PyQt6 import uic
from docx import Document

class TextEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('text_editor.ui', self)

        self.opened_file = None

        self.pushButton.clicked.connect(self.open_file)
        self.pushButton_2.clicked.connect(self.save_file)
        self.pushButton_3.clicked.connect(self.save_file_as)

        self.textBrowser.setStyleSheet("border: none; background-color: transparent;")
        self.textBrowser.setHtml('''
            <h1>Welcome to My Text Editor!</h1>
        ''')

        self.fontComboBox.currentFontChanged.connect(self.change_font)
        self.comboBox_2.setStyleSheet("background-color: white")
        self.comboBox_2.addItems(["8", "10", "12", "14", "16", "18", "20"])
        self.comboBox_2.currentIndexChanged.connect(self.change_font_size)

        self.pushButton_4.clicked.connect(self.close_file)
        self.pushButton_5.clicked.connect(self.choose_text_color)
        self.pushButton_6.clicked.connect(lambda: self.set_bold(self.textEdit.fontWeight() != QFont.Weight.Bold))
        self.pushButton_7.clicked.connect(lambda: self.set_italic(not self.textEdit.fontItalic()))
        self.pushButton_8.clicked.connect(lambda: self.set_underline(not self.textEdit.fontUnderline()))

    def open_file(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "Open File", "", "Text Files (*.txt);;HTML Files (*.html);;Word Files (*.docx)")
        if file_name:
            if file_name.endswith('.txt'):
                with open(file_name, 'r', encoding='utf-8') as file:
                    text = file.read()
                    self.textEdit.setPlainText(text)
            elif file_name.endswith('.html'):
                with open(file_name, 'r', encoding='utf-8') as file:
                    text = file.read()
                    self.textEdit.setHtml(text)
            elif file_name.endswith('.docx'):
                doc = Document(file_name)
                text = '\n'.join([para.text for para in doc.paragraphs])
                self.textEdit.setPlainText(text)
            self.opened_file = file_name

    def save_file(self):
        if self.opened_file:
            if self.opened_file.endswith('.txt'):
                self.save_as_txt(self.opened_file)
            elif self.opened_file.endswith('.html'):
                self.save_as_html(self.opened_file)
            elif self.opened_file.endswith('.docx'):
                self.save_as_docx(self.opened_file)
            QMessageBox.information(self, "Saved", "File successfully saved.")
        else:
            QMessageBox.information(self, "System Notification", "File does not exist, you need to create it.")
            self.save_file_as()

    def save_file_as(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Text Files (*.txt);;HTML Files (*.html);;Word Files (*.docx)")
        if file_name:
            if file_name.endswith('.txt'):
                self.save_as_txt(file_name)
            elif file_name.endswith('.html'):
                self.save_as_html(file_name)
            elif file_name.endswith('.docx'):
                self.save_as_docx(file_name)
            QMessageBox.information(self, "Saved", "File successfully saved.")
            self.textEdit.clear()
            self.opened_file = file_name

    def save_as_txt(self, file_name):
        with open(file_name, 'w', encoding='utf-8') as file:
            text = self.textEdit.toPlainText()
            file.write(text)

    def save_as_html(self, file_name):
        with open(file_name, 'w', encoding='utf-8') as file:
            text = self.textEdit.toHtml()
            file.write(text)

    def save_as_docx(self, file_name):
        doc = Document()
        text = self.textEdit.toPlainText()
        doc.add_paragraph(text)
        doc.save(file_name)

    def close_file(self):
        self.textEdit.clear()
        self.opened_file = None

    def change_font(self, font):
        cursor = self.textEdit.textCursor()
        format = QTextCharFormat()
        format.setFont(font)
        cursor.mergeCharFormat(format)

    def change_font_size(self):
        font_size = int(self.sender().currentText())
        cursor = self.textEdit.textCursor()
        format = QTextCharFormat()
        format.setFontPointSize(font_size)
        cursor.mergeCharFormat(format)

    def choose_text_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            self.set_text_color(color)

    def set_text_color(self, color):
        cursor = self.textEdit.textCursor()
        format = QTextCharFormat()
        format.setForeground(QColor(color))
        cursor.mergeCharFormat(format)

    def set_bold(self, bold):
        cursor = self.textEdit.textCursor()
        format = QTextCharFormat()
        if bold:
            format.setFontWeight(QFont.Weight.Bold)
        else:
            format.setFontWeight(QFont.Weight.Normal)
        cursor.mergeCharFormat(format)

    def set_italic(self, italic):
        cursor = self.textEdit.textCursor()
        format = QTextCharFormat()
        format.setFontItalic(italic)
        cursor.mergeCharFormat(format)

    def set_underline(self, underline):
        cursor = self.textEdit.textCursor()
        format = QTextCharFormat()
        format.setFontUnderline(underline)
        cursor.mergeCharFormat(format)

if __name__ == '__main__':
    app = QApplication([])
    editor = TextEditor()
    editor.show()
    sys.exit(app.exec())
